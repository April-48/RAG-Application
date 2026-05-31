"""Extract plain text from DOCX files.

DOCX files do not have stable page numbers like PDFs, so I return everything
as one big PageText block with page_number=None. Downstream chunking still
works — you just cannot do page-specific lookups on DOCX uploads.

Legacy .doc (Word 97) is not supported. Only modern .docx (ZIP + XML) works.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.table import Table

from app.rag.pdf_parser import PageText


def _table_lines(table: Table) -> list[str]:
    """Turn a DOCX table into plain text lines.

    Each row becomes one line: cell1 | cell2 | cell3.
    Empty cells are skipped so the line does not end with stray pipes.
    """
    lines: list[str] = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        cells = [cell for cell in cells if cell]
        if cells:
            lines.append(" | ".join(cells))
    return lines


def extract_docx_pages(path: str | Path) -> list[PageText]:
    """Read a DOCX file and return its text as a single PageText.

    I walk paragraphs first, then tables, and join everything with newlines.
    If the file is empty I still return one PageText with empty text so the
    ingestion pipeline can mark the document as failed (zero chunks).
    """
    document = DocxDocument(str(path))

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        parts.extend(_table_lines(table))

    return [PageText(page_number=None, text="\n".join(parts))]
